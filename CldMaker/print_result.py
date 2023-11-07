import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def print_result_docx(Title: str, IMAGE_DIR, doc_name, results_df):
    doc = Document()  # create doc
    doc.add_heading(Title, 0)  # add heading

    # create table with two rows and columns (Per row images are 3)
    table = doc.add_table(rows=0, cols=3, style="Table Grid")
    images = os.listdir(IMAGE_DIR)

    table.autofit = True
    table.allow_autofit = True
    # table.allow_autofit
    for i in results_df.index:
        image_row = table.add_row()  # add row to table for images
        # cap_row = table.add_row() # add row to table for image name

        image_name = images.pop()
        image_row.cells[0].text = "Case " + str(i)

        # add image to table
        # set_cell_margins(image_row.cells[1], top=100, start=100, bottom=100, end=50)
        # add image to cell and align center
        paragraph = image_row.cells[1].paragraphs[0]
        paragraph.add_run().add_picture(f"{IMAGE_DIR}/label_{i}.png", width=Inches(2.5))

        # set_cell_margins(image_row.cells[1], top=100, start=100, bottom=100, end=50)
        paragraph = image_row.cells[2].paragraphs[0]
        paragraph.add_run().add_picture(f"{IMAGE_DIR}/predict_{i}.png", width=Inches(2.5))

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # save to file
    doc.save(f"{doc_name}.docx")
